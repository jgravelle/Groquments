import json
import os
import sys
import re
from pocketgroq import GroqProvider
from typing import Dict, Any
from docx import Document
import tkinter as tk
from tkinter import filedialog
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

class Groquments:
    def __init__(self):
        self.groq_provider = GroqProvider(api_key=os.getenv("GROQ_API_KEY"))
        self.model = os.getenv("GROQ_MODEL", "llama3-8b-8192")
        self.max_tokens = int(os.getenv("MAX_TOKENS", "1024"))
        self.temperature = float(os.getenv("TEMPERATURE", "0.7"))
        self.source_document = {}
        self.target_document = {}
        self.input_dir = os.getenv("INPUT_DIR", "input_documents")
        self.output_dir = os.getenv("OUTPUT_DIR", "output_documents")

    def read_docx(self, file_path: str) -> Dict[str, str]:
        doc = Document(file_path)
        document_content = {}
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if ':' in text:
                label, value = text.split(':', 1)
                document_content[label.strip()] = value.strip()
            elif text and not text.endswith(':'):
                document_content[text] = ""
        return document_content

    def create_js_representation(self) -> str:
        js_object = json.dumps(self.source_document, indent=2)
        return f"const sourceDocument = {js_object};"

    def display_js_representation(self):
        print("\nJavaScript representation of the source document:")
        print(self.create_js_representation())

    def extract_json_from_text(self, text: str) -> str:
        json_match = re.search(r'\{[\s\S]*\}', text)
        return json_match.group(0) if json_match else "{}"

    def ai_field_mapping(self) -> Dict[str, str]:
        prompt = f"""
        Given the following source document structure:
        {json.dumps(self.source_document, indent=2)}

        And the following target document structure:
        {json.dumps(self.target_document, indent=2)}

        Please map the fields from the source document to the target document.
        Return your answer as a JSON object where the keys are the target document fields
        and the values are the corresponding source document fields.
        If there's no clear match, use an empty string as the value.
        Only return the JSON object, without any additional explanation.
        """

        response = self.groq_provider.generate(
            prompt,
            model=self.model,
            max_tokens=self.max_tokens,
            temperature=self.temperature
        )

        print("AI response:", response)  # Print the full AI response for debugging

        json_str = self.extract_json_from_text(response)
        
        try:
            field_mapping = json.loads(json_str)
            return field_mapping
        except json.JSONDecodeError:
            print("Error: Unable to parse AI response. Using fallback mapping.")
            return {key: "" for key in self.target_document.keys()}

    def fill_target_document(self):
        field_mapping = self.ai_field_mapping()
        
        for target_key, source_key in field_mapping.items():
            if source_key in self.source_document:
                self.target_document[target_key] = self.source_document[source_key]
            else:
                self.target_document[target_key] = ""

    def save_filled_document(self, output_path: str):
        doc = Document()
        for key, value in self.target_document.items():
            doc.add_paragraph(f"{key}: {value}")
        doc.save(output_path)

    def run(self):
        print("Welcome to Groquments!")
        
        root = tk.Tk()
        root.withdraw()

        source_file = filedialog.askopenfilename(title="Select source .docx file", 
                                                 filetypes=[("Word Document", "*.docx")],
                                                 initialdir=self.input_dir)
        if not source_file:
            print("No source file selected. Exiting.")
            sys.exit()

        target_file = filedialog.askopenfilename(title="Select target .docx file", 
                                                 filetypes=[("Word Document", "*.docx")],
                                                 initialdir=self.input_dir)
        if not target_file:
            print("No target file selected. Exiting.")
            sys.exit()

        self.source_document = self.read_docx(source_file)
        self.display_js_representation()
        
        self.target_document = self.read_docx(target_file)
        
        print("\nFilling out the target document...")
        self.fill_target_document()
        
        print("\nFilled target document:")
        print(json.dumps(self.target_document, indent=2))

        os.makedirs(self.output_dir, exist_ok=True)
        output_file = os.path.join(self.output_dir, "filled_" + os.path.basename(target_file))
        self.save_filled_document(output_file)
        print(f"\nFilled document saved as: {output_file}")

if __name__ == "__main__":
    groquments = Groquments()
    groquments.run()